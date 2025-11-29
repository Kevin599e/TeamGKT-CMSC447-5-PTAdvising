from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from collections import OrderedDict, defaultdict
from pathlib import Path
import json

def _shade_cell(cell, fill_hex: str = "C6EFCE"):
    """
    Apply background shading to a cell.
    fill_hex: hex color without '#', e.g. 'C6EFCE' (light green).
    """
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill_hex)
    tc_pr.append(shd)


def render_table(doc, table_json_str):
    try:
        data = json.loads(table_json_str)
    except Exception:
        doc.add_paragraph("[Could not parse table]")
        return

    columns = data.get("columns", [])
    rows = data.get("rows", [])

    if not columns or not rows:
        doc.add_paragraph("[No table data]")
        return

    # We will use all columns *except* the first ("Term") as table columns
    # and show the term as a heading above each table.
    visible_cols = columns[1:] if len(columns) > 1 else columns

    # Group rows by term, and track benchmark rows to show as notes
    grouped = OrderedDict()          # term -> list of row_values (without term)
    benchmarks = defaultdict(list)   # term -> list of benchmark strings
    last_term = None

    for row in rows:
        if not row:
            continue
        term = str(row[0])

        # Detect benchmark rows by name
        if "benchmark" in term.lower():
            # Attach this benchmark text to the last non-benchmark term we saw
            note = row[3] if len(row) > 3 else ""
            if last_term is not None and note:
                benchmarks[last_term].append(note)
            continue

        # Normal term row: store without the term column
        last_term = term
        grouped.setdefault(term, []).append(row[1:])

    # Figure out which index is "Credits" so we can center it
    try:
        credits_idx = visible_cols.index("Credits")
    except ValueError:
        credits_idx = None

    # ----- Render a separate table for each term -----
    first_term = True
    for term, term_rows in grouped.items():
        if not first_term:
            # Add a blank paragraph between terms for spacing
            doc.add_paragraph()
        first_term = False

        # Term heading (e.g., "Year 1 - Fall")
        heading_para = doc.add_paragraph(term)
        heading_run = heading_para.runs[0]
        heading_run.bold = True

        # Create table: 1 header row + data rows
        table = doc.add_table(rows=1 + len(term_rows), cols=len(visible_cols))
        table.style = "Table Grid"  # grid lines between rows and columns

        # Header row
        hdr_cells = table.rows[0].cells
        for j, col_name in enumerate(visible_cols):
            if j >= len(hdr_cells):
                break
            cell = hdr_cells[j]
            cell.text = str(col_name)

            # Bold + shaded header
            for p in cell.paragraphs:
                for run in p.runs:
                    run.bold = True
            _shade_cell(cell, "C6EFCE")

            # Center the Credits header
            if str(col_name).lower() == "credits":
                for p in cell.paragraphs:
                    p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

        # Body rows
        for i, row_vals in enumerate(term_rows):
            cells = table.rows[i + 1].cells
            for j, cell_val in enumerate(row_vals):
                if j >= len(cells):
                    break
                cell = cells[j]
                cell.text = str(cell_val)

                # Center numeric credits like in sample plan
                if credits_idx is not None and j == credits_idx:
                    for p in cell.paragraphs:
                        p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

        # Benchmarks attached to this term (if any)
        if term in benchmarks:
            for note in benchmarks[term]:
                p = doc.add_paragraph()
                r_label = p.add_run("Benchmarks: ")
                r_label.bold = True
                p.add_run(note)


def render_packet_docx(packet, sections, export_dir="exports"):
    export_path = Path(export_dir)
    export_path.mkdir(parents=True, exist_ok=True)
    filename = export_path / f"packet_{packet.id}.docx"

    doc = Document()

    # Header info
    doc.add_heading("UMBC Advising Packet", level=1)
    doc.add_paragraph(
        f"Student: {packet.request.student_name} <{packet.request.student_email}>"
    )
    doc.add_paragraph(
        f"Source Institution: {packet.request.source_institution or '-'}"
    )
    doc.add_paragraph(
        f"Target Program: {packet.request.target_program or '-'}"
    )
    doc.add_paragraph("")

    for s in sections:
        doc.add_heading(s.title, level=2)

        if s.content_type in ("table", "audit_table"):
            render_table(doc, s.content or "{}")
        else:
            # assume plain text / markdown-ish
            doc.add_paragraph(s.content or "")

    doc.save(str(filename))
    return str(filename)
